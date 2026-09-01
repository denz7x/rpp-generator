"""
Backend logic: generation via Gemini untuk MODUL AJAR KURIKULUM MERDEKA
(Deep Learning), mengikuti struktur dokumen contoh resmi:

  - Cover (tabel)
  - MODUL AJAR DEEP LEARNING / MATA PELAJARAN / BAB
    A. IDENTITAS MODUL
    B. IDENTIFIKASI KESIAPAN PESERTA DIDIK
    C. KARAKTERISTIK MATERI PELAJARAN
    D. DIMENSI PROFIL LULUSAN
  DESAIN PEMBELAJARAN
    A. CAPAIAN PEMBELAJARAN (CP)
    B. LINTAS DISIPLIN ILMU
    C. TUJUAN PEMBELAJARAN
    D. TOPIK PEMBELAJARAN KONTEKSTUAL
    E. KERANGKA PEMBELAJARAN
    F. LANGKAH-LANGKAH PEMBELAJARAN BERDIFERENSIASI (per pertemuan)
    G. ASESMEN PEMBELAJARAN
  - Tanda tangan (tabel)

Sistem ekstraksi JSON tetap sama (anti-gagal + error handling detail).
"""

import io
import json
import streamlit as st

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_helpers import (
    style_run, add_paragraph_in_cell, add_bottom_border,
    add_bold_lead_bullet, add_plain_bullet, add_manual_numbered,
    add_body_paragraph, set_heading_font, set_table_borders, set_col_widths,
)

try:
    import google.generativeai as genai
except ImportError:
    genai = None

# 8 Dimensi Profil Lulusan (Deep Learning) — nama baku, tidak diubah oleh AI.
DIMENSI_PROFIL_LULUSAN_NAMA = [
    "Keimanan dan Ketakwaan terhadap Tuhan Yang Maha Esa, dan Berakhlak Mulia",
    "Kewargaan",
    "Penalaran Kritis",
    "Kreativitas",
    "Kolaborasi",
    "Kemandirian",
    "Kesehatan",
    "Komunikasi",
]

# ==========================================================================
# 1. GENERASI KONTEN VIA GEMINI (Sistem Anti-Gagal)
# ==========================================================================

def _clean_json(text):
    """
    Ekstraktor super aman: Memaksa mencari kurung kurawal pertama {
    dan kurung kurawal terakhir } untuk mengabaikan teks Markdown tambahan dari AI.
    """
    text = text.strip()
    start = text.find('{')
    end = text.rfind('}')

    if start != -1 and end != -1:
        return text[start:end + 1]
    return text


def _ctx_block(ctx):
    return f"""
Mata Pelajaran : {ctx['mapel']}
Jenjang/Fase   : {ctx['jenjang']} / Fase {ctx['fase']}
Kelas/Semester : {ctx['kelas']} / {ctx['semester']}
BAB / Tema     : {ctx['tema']}
Materi/Topik   : {ctx['materi']}
Alokasi Waktu Total : {ctx['waktu']}
Jumlah Blok Pertemuan yang harus dibuat : {ctx['jumlah_pertemuan']}
Tahun Pelajaran : {ctx.get('tahun_pelajaran', '')}
Model Pembelajaran pilihan guru (acuan utama) : {ctx['model_pembelajaran']}
Daftar Profil Pelajar Pancasila yang dituju (dipakai untuk "Integrasi Nilai dan Karakter") : {', '.join(ctx['profil'])}
Dimensi Profil Lulusan yang WAJIB dijelaskan (urutan tetap, jangan diubah namanya), sebanyak {len(DIMENSI_PROFIL_LULUSAN_NAMA)} item:
{chr(10).join(f"{i+1}. {n}" for i, n in enumerate(DIMENSI_PROFIL_LULUSAN_NAMA))}
""".strip()


PROMPT_BAGIAN_UMUM = """
Kamu adalah asisten ahli penyusun MODUL AJAR KURIKULUM MERDEKA dengan pendekatan
DEEP LEARNING (Mindful, Meaningful, Joyful Learning) untuk guru di Indonesia.
Buatkan bagian IDENTITAS MODUL s.d. DESAIN PEMBELAJARAN (kecuali langkah per
pertemuan dan asesmen) untuk konteks berikut:

{ctx}

Tulis dengan bahasa Indonesia baku, kontekstual sesuai materi di atas (jangan
generik), dan konsisten dengan jumlah blok pertemuan yang diminta.

Keluarkan HANYA JSON murni dengan skema PERSIS berikut (semua field wajib diisi):

{{
  "identifikasi_pengetahuan_awal": "1-2 kalimat pengetahuan awal peserta didik terkait materi",
  "identifikasi_minat": "1-2 kalimat minat peserta didik yang relevan",
  "identifikasi_latar_belakang": "1-2 kalimat latar belakang peserta didik yang relevan",
  "kebutuhan_visual": "kalimat kebutuhan belajar visual terkait materi",
  "kebutuhan_auditori": "kalimat kebutuhan belajar auditori terkait materi",
  "kebutuhan_kinestetik": "kalimat kebutuhan belajar kinestetik terkait materi",
  "karakteristik_konseptual": "kalimat konsep-konsep kunci yang akan dipahami peserta didik",
  "karakteristik_prosedural": "kalimat keterampilan prosedural yang akan dikuasai peserta didik",
  "relevansi_kehidupan_nyata": "1-2 kalimat relevansi materi dengan kehidupan nyata peserta didik",
  "tingkat_kesulitan": "1 kalimat tingkat kesulitan materi (mudah/sedang/sulit) beserta alasan singkat",
  "struktur_materi": "1-2 kalimat bagaimana materi disusun secara sistematis",
  "integrasi_nilai": [
    {{"profil": "<harus persis salah satu dari daftar Profil Pelajar Pancasila di konteks>", "deskripsi": "kalimat penjelasan integrasi nilai ini dalam materi"}}
  ],
  "dimensi_profil_lulusan_desc": [
    "deskripsi untuk dimensi ke-1 sesuai urutan daftar Dimensi Profil Lulusan di konteks",
    "deskripsi untuk dimensi ke-2", "... dst sampai tepat 8 item, urutan harus sama persis dengan daftar di konteks"
  ],
  "capaian_pemahaman_konsep": "1 paragraf capaian pembelajaran aspek pemahaman konsep, kontekstual dengan materi",
  "capaian_keterampilan_proses": "1 paragraf capaian pembelajaran aspek keterampilan proses, kontekstual dengan materi",
  "lintas_disiplin": [
    {{"disiplin": "nama mata pelajaran/disiplin ilmu terkait", "deskripsi": "kalimat singkat keterkaitannya dengan materi"}}
  ],
  "tujuan_pembelajaran": [
    {{"pertemuan": "1-2", "jp": 4, "deskripsi": "kalimat tujuan pembelajaran spesifik untuk blok pertemuan ini"}}
  ],
  "topik_kontekstual": "1-2 kalimat judul & deskripsi topik pembelajaran kontekstual untuk keseluruhan bab",
  "pendekatan_mindful": "1-2 kalimat penerapan Mindful Learning pada materi ini",
  "pendekatan_meaningful": "1-2 kalimat penerapan Meaningful Learning pada materi ini",
  "pendekatan_joyful": "1-2 kalimat penerapan Joyful Learning pada materi ini",
  "metode_pembelajaran": "daftar metode pembelajaran dipisah koma, kontekstual dengan materi",
  "diferensiasi_konten": "1-2 kalimat strategi diferensiasi konten",
  "diferensiasi_proses": "1-2 kalimat strategi diferensiasi proses",
  "diferensiasi_produk": "1-2 kalimat strategi diferensiasi produk",
  "kemitraan_sekolah": "1-2 kalimat kemitraan dengan lingkungan sekolah terkait materi",
  "kemitraan_luar_sekolah": "1-2 kalimat kemitraan dengan lingkungan luar sekolah/masyarakat terkait materi",
  "kemitraan_digital": "1-2 kalimat mitra digital/platform yang relevan",
  "lingkungan_ruang_fisik": "1-2 kalimat pengaturan ruang fisik yang relevan",
  "lingkungan_ruang_virtual": "1-2 kalimat pemanfaatan ruang virtual yang relevan",
  "lingkungan_budaya_belajar": "1-2 kalimat budaya belajar yang ingin dibangun",
  "pemanfaatan_digital": [
    "poin pemanfaatan sumber/media digital 1", "poin 2", "poin 3"
  ]
}}

PENTING: array "tujuan_pembelajaran" jumlah itemnya HARUS SAMA PERSIS dengan
"Jumlah Blok Pertemuan yang harus dibuat" pada konteks, dan totalnya (jumlah
"jp" semua item) harus masuk akal dibandingkan "Alokasi Waktu Total". Array
"dimensi_profil_lulusan_desc" HARUS berjumlah tepat 8 item.
"""

PROMPT_LANGKAH_ASESMEN = """
Kamu adalah asisten ahli penyusun MODUL AJAR KURIKULUM MERDEKA dengan pendekatan
DEEP LEARNING (Mindful, Meaningful, Joyful Learning) untuk guru di Indonesia.
Buatkan bagian "F. LANGKAH-LANGKAH PEMBELAJARAN BERDIFERENSIASI" (rinci per
blok pertemuan, meniru gaya RPP/modul ajar resmi Kemendikbud dengan kegiatan
pendahuluan/inti/penutup beserta alokasi menit) dan bagian
"G. ASESMEN PEMBELAJARAN" untuk konteks berikut:

{ctx}

Keluarkan HANYA JSON murni dengan skema PERSIS berikut:

{{
  "langkah_pembelajaran": [
    {{
      "pertemuan": "1-2",
      "jp_label": "4 JP : 2 x 80 Menit",
      "topik": "JUDUL TOPIK PERTEMUAN INI (huruf kapital, singkat)",
      "pendahuluan_menit": "15 MENIT",
      "pendahuluan": [
        "Salam dan Doa: kalimat kegiatan",
        "Presensi: kalimat kegiatan",
        "Apersepsi: kalimat kegiatan kontekstual dengan materi pertemuan ini",
        "Tujuan Pembelajaran: kalimat penyampaian tujuan"
      ],
      "inti_menit": "55 MENIT",
      "inti": [
        "Eksplorasi Konsep: kalimat kegiatan sesuai model pembelajaran",
        "Aktivitas (individu/kelompok): kalimat kegiatan",
        "kegiatan inti lain jika relevan"
      ],
      "penutup_menit": "10 MENIT",
      "penutup": [
        "Refleksi: kalimat refleksi",
        "Tindak Lanjut: kalimat tindak lanjut",
        "Penutup: Salam dan doa."
      ]
    }}
  ],
  "asesmen_diagnostik": [
    "poin teknik asesmen diagnostik 1 (jelaskan tekniknya dan contoh pertanyaannya, kontekstual dengan materi)",
    "poin 2"
  ],
  "asesmen_formatif": [
    "poin teknik asesmen formatif 1 (kontekstual dengan materi)",
    "poin 2", "poin 3"
  ],
  "asesmen_sumatif": [
    "poin teknik asesmen sumatif 1 (kontekstual dengan materi, sebutkan bentuk produk/praktik/tes)",
    "poin 2"
  ],
  "soal_pg": [
    {{"soal": "teks soal pilihan ganda kontekstual dengan materi", "a": "pilihan a", "b": "pilihan b", "c": "pilihan c", "d": "pilihan d", "kunci": "b"}}
  ],
  "soal_esai": [
    "teks soal esai 1 kontekstual dengan materi",
    "teks soal esai 2"
  ]
}}

PENTING: array "langkah_pembelajaran" jumlah itemnya HARUS SAMA PERSIS dengan
"Jumlah Blok Pertemuan yang harus dibuat" pada konteks, dan gunakan label
"pertemuan" yang berurutan dan tidak tumpang tindih (misal "1-2", lalu "3",
lalu "4-5", dst) sehingga totalnya konsisten dengan "Alokasi Waktu Total".
Buat tepat 3 soal pilihan ganda dan tepat 3 soal esai.
"""


def _call_gemini_json(model_name, prompt, tahap_nama):
    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)

        text_bersih = _clean_json(response.text)
        return json.loads(text_bersih)

    except json.JSONDecodeError as je:
        st.error(f"❌ [Tahap {tahap_nama}] AI gagal memberikan format yang tepat.")
        if 'response' in locals() and hasattr(response, 'text'):
            with st.expander(f"🔍 Klik untuk lihat balasan AI yang error ({tahap_nama})"):
                st.code(response.text)
        return None
    except Exception as e:
        st.error(f"⚠️ [Tahap {tahap_nama}] Error Sistem: {str(e)}")
        return None


def generate_bagian_umum(model_name, ctx):
    prompt = PROMPT_BAGIAN_UMUM.format(ctx=_ctx_block(ctx))
    return _call_gemini_json(model_name, prompt, "Identitas & Desain Pembelajaran")


def generate_asesmen_lampiran(model_name, ctx):
    prompt = PROMPT_LANGKAH_ASESMEN.format(ctx=_ctx_block(ctx))
    return _call_gemini_json(model_name, prompt, "Langkah Pembelajaran & Asesmen")


# ==========================================================================
# 2. PEMBUATAN DOCX
# ==========================================================================

def _g(d, key, default):
    """Ambil field dari dict AI dengan fallback aman."""
    if not d:
        return default
    val = d.get(key, default)
    return val if val not in (None, "") else default


def _heading(doc, text, level, align=None, italic=False):
    h = doc.add_heading(text, level=level)
    if align:
        h.alignment = align
    set_heading_font(h)
    if italic:
        for run in h.runs:
            run.italic = True
    return h


def create_docx(data_input, ai_umum, ai_asesmen):
    ai_umum = ai_umum or {}
    ai_asesmen = ai_asesmen or {}

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    bab_judul = data_input['tema']
    mapel_upper = data_input['mapel'].upper()

    # ---------------- COVER (tabel — format dipertahankan) ----------------
    for _ in range(4):
        doc.add_paragraph()

    cover_tbl = doc.add_table(rows=5, cols=1)
    cover_tbl.autofit = True
    p = cover_tbl.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("MODUL AJAR"), bold=True, size=16)
    p2 = cover_tbl.rows[1].cells[0].add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("KURIKULUM MERDEKA (Deep Learning)")
    style_run(r2, bold=True, italic=True, size=14)

    id_labels = [
        ("Nama Sekolah", data_input['sekolah']),
        ("Nama Penyusun", data_input['guru']),
        ("NIP", data_input.get('nik', '')),
        ("Mata Pelajaran", data_input['mapel']),
        ("Fase " + str(data_input['fase']) + ", Kelas / Semester",
         f"{data_input['kelas']} / {data_input['semester']}"),
    ]
    row2 = cover_tbl.rows[2].cells[0]
    row2.paragraphs[0].text = ""
    first = True
    for label, val in id_labels:
        shown = val if val else "…………………………"
        add_paragraph_in_cell(row2, f"{label}\t:\t{shown}", bold=True, first=first)
        first = False
    cover_tbl.rows[3].cells[0].paragraphs[0].text = ""
    cover_tbl.rows[4].cells[0].paragraphs[0].text = ""
    set_table_borders(cover_tbl, size=6)

    doc.add_page_break()

    # ---------------- JUDUL HALAMAN ISI ----------------
    _heading(doc, f"MODUL AJAR DEEP LEARNING", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _heading(doc, f"MATA PELAJARAN : {mapel_upper}", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _heading(doc, f"BAB {data_input.get('nomor_bab', '1')}: {bab_judul.upper()}", 2,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    # ============ A. IDENTITAS MODUL ============
    _heading(doc, "A. IDENTITAS MODUL", 3)
    identitas_pairs = [
        ("Nama Sekolah", data_input['sekolah']),
        ("Nama Penyusun", data_input['guru']),
        ("Mata Pelajaran", data_input['mapel']),
        ("Kelas / Fase /Semester",
         f"{data_input['kelas']} / {data_input['fase']} / {data_input['semester']}"),
        ("Alokasi Waktu", data_input['waktu']),
        ("Tahun Pelajaran", data_input.get('tahun_pelajaran', '')),
    ]
    for label, val in identitas_pairs:
        shown = val if val else "…………………"
        add_body_paragraph(doc, f"{label} : {shown}", bold=True, space_after=2)

    # ============ B. IDENTIFIKASI KESIAPAN PESERTA DIDIK ============
    _heading(doc, "B. IDENTIFIKASI KESIAPAN PESERTA DIDIK", 3)
    add_bold_lead_bullet(doc, "Pengetahuan Awal", _g(ai_umum, 'identifikasi_pengetahuan_awal', '-'))
    add_bold_lead_bullet(doc, "Minat", _g(ai_umum, 'identifikasi_minat', '-'))
    add_bold_lead_bullet(doc, "Latar Belakang", _g(ai_umum, 'identifikasi_latar_belakang', '-'))
    add_bold_lead_bullet(doc, "Kebutuhan Belajar", "")
    add_bold_lead_bullet(doc, "Visual", _g(ai_umum, 'kebutuhan_visual', '-'), level=1)
    add_bold_lead_bullet(doc, "Auditori", _g(ai_umum, 'kebutuhan_auditori', '-'), level=1)
    add_bold_lead_bullet(doc, "Kinestetik", _g(ai_umum, 'kebutuhan_kinestetik', '-'), level=1)

    # ============ C. KARAKTERISTIK MATERI PELAJARAN ============
    _heading(doc, "C. KARAKTERISTIK MATERI PELAJARAN", 3)
    add_bold_lead_bullet(doc, "Jenis Pengetahuan yang Akan Dicapai", "")
    add_bold_lead_bullet(doc, "Konseptual", _g(ai_umum, 'karakteristik_konseptual', '-'), level=1)
    add_bold_lead_bullet(doc, "Prosedural", _g(ai_umum, 'karakteristik_prosedural', '-'), level=1)
    add_bold_lead_bullet(doc, "Relevansi dengan Kehidupan Nyata Peserta Didik",
                          _g(ai_umum, 'relevansi_kehidupan_nyata', '-'))
    add_bold_lead_bullet(doc, "Tingkat Kesulitan", _g(ai_umum, 'tingkat_kesulitan', '-'))
    add_bold_lead_bullet(doc, "Struktur Materi", _g(ai_umum, 'struktur_materi', '-'))
    add_bold_lead_bullet(doc, "Integrasi Nilai dan Karakter", "")
    integrasi = _g(ai_umum, 'integrasi_nilai', [])
    if integrasi:
        for item in integrasi:
            add_bold_lead_bullet(doc, item.get('profil', '-'), item.get('deskripsi', '-'), level=1)
    else:
        for profil in data_input['profil']:
            add_bold_lead_bullet(doc, profil, "-", level=1)

    # ============ D. DIMENSI PROFIL LULUSAN ============
    _heading(doc, "D. DIMENSI PROFIL LULUSAN", 3)
    dpl_desc = _g(ai_umum, 'dimensi_profil_lulusan_desc', [])
    for i, nama in enumerate(DIMENSI_PROFIL_LULUSAN_NAMA):
        desc = dpl_desc[i] if i < len(dpl_desc) else "-"
        add_bold_lead_bullet(doc, nama, desc)

    # ============ DESAIN PEMBELAJARAN ============
    _heading(doc, "DESAIN PEMBELAJARAN", 2)

    _heading(doc, "A. CAPAIAN PEMBELAJARAN (CP)", 3)
    add_body_paragraph(doc, f"Pada akhir Fase {data_input['fase']}, murid memiliki kemampuan sebagai berikut.")
    add_bold_lead_bullet(doc, "Pemahaman Konsep", _g(ai_umum, 'capaian_pemahaman_konsep', '-'))
    add_bold_lead_bullet(doc, "Keterampilan Proses", _g(ai_umum, 'capaian_keterampilan_proses', '-'))

    _heading(doc, "B. LINTAS DISIPLIN ILMU", 3)
    for item in _g(ai_umum, 'lintas_disiplin', []):
        add_bold_lead_bullet(doc, item.get('disiplin', '-'), item.get('deskripsi', '-'))

    _heading(doc, "C. TUJUAN PEMBELAJARAN", 3)
    for item in _g(ai_umum, 'tujuan_pembelajaran', []):
        add_bold_lead_bullet(
            doc, f"Pertemuan {item.get('pertemuan', '-')}",
            f"{item.get('deskripsi', '-')} ({item.get('jp', '-')} JP)"
        )

    _heading(doc, "D. TOPIK PEMBELAJARAN KONTEKSTUAL", 3)
    add_body_paragraph(doc, f"{bab_judul.upper()}: {_g(ai_umum, 'topik_kontekstual', '-')}")

    _heading(doc, "E. KERANGKA PEMBELAJARAN", 3)
    _heading(doc, "PRAKTIK PEDAGOGIK", 4)
    add_bold_lead_bullet(doc, "Model Pembelajaran", data_input['model_pembelajaran'])
    add_bold_lead_bullet(doc, "Pendekatan", "Deep Learning (Mindful, Meaningful, Joyful Learning)")
    add_bold_lead_bullet(doc, "Mindful Learning", _g(ai_umum, 'pendekatan_mindful', '-'), level=1)
    add_bold_lead_bullet(doc, "Meaningful Learning", _g(ai_umum, 'pendekatan_meaningful', '-'), level=1)
    add_bold_lead_bullet(doc, "Joyful Learning", _g(ai_umum, 'pendekatan_joyful', '-'), level=1)
    add_bold_lead_bullet(doc, "Metode Pembelajaran", _g(ai_umum, 'metode_pembelajaran', '-'))
    add_bold_lead_bullet(doc, "Strategi Pembelajaran Berdiferensiasi", "")
    add_bold_lead_bullet(doc, "Diferensiasi Konten", _g(ai_umum, 'diferensiasi_konten', '-'), level=1)
    add_bold_lead_bullet(doc, "Diferensiasi Proses", _g(ai_umum, 'diferensiasi_proses', '-'), level=1)
    add_bold_lead_bullet(doc, "Diferensiasi Produk", _g(ai_umum, 'diferensiasi_produk', '-'), level=1)

    _heading(doc, "KEMITRAAN PEMBELAJARAN", 4)
    add_bold_lead_bullet(doc, "Lingkungan Sekolah", _g(ai_umum, 'kemitraan_sekolah', '-'))
    add_bold_lead_bullet(doc, "Lingkungan Luar Sekolah/Masyarakat", _g(ai_umum, 'kemitraan_luar_sekolah', '-'))
    add_bold_lead_bullet(doc, "Mitra Digital", _g(ai_umum, 'kemitraan_digital', '-'))

    _heading(doc, "LINGKUNGAN BELAJAR", 4)
    add_bold_lead_bullet(doc, "Ruang Fisik", _g(ai_umum, 'lingkungan_ruang_fisik', '-'))
    add_bold_lead_bullet(doc, "Ruang Virtual", _g(ai_umum, 'lingkungan_ruang_virtual', '-'))
    add_bold_lead_bullet(doc, "Budaya Belajar", _g(ai_umum, 'lingkungan_budaya_belajar', '-'))

    _heading(doc, "PEMANFAATAN DIGITAL", 4)
    for poin in _g(ai_umum, 'pemanfaatan_digital', ['-']):
        add_plain_bullet(doc, poin)

    # ============ F. LANGKAH-LANGKAH PEMBELAJARAN BERDIFERENSIASI ============
    _heading(doc, "F. LANGKAH-LANGKAH PEMBELAJARAN BERDIFERENSIASI", 3)
    for pert in _g(ai_asesmen, 'langkah_pembelajaran', []):
        _heading(doc, f"PERTEMUAN {pert.get('pertemuan', '-')} ({pert.get('jp_label', '-')})", 4)
        add_bold_lead_bullet(doc, "Topik", pert.get('topik', '-'))
        add_bold_lead_bullet(doc, f"KEGIATAN PENDAHULUAN ({pert.get('pendahuluan_menit', '-')})", "")
        for item in pert.get('pendahuluan', ['-']):
            add_plain_bullet(doc, item, level=1)
        add_bold_lead_bullet(doc, f"KEGIATAN INTI ({pert.get('inti_menit', '-')})", "")
        for item in pert.get('inti', ['-']):
            add_plain_bullet(doc, item, level=1)
        add_bold_lead_bullet(doc, f"KEGIATAN PENUTUP ({pert.get('penutup_menit', '-')})", "")
        for item in pert.get('penutup', ['-']):
            add_plain_bullet(doc, item, level=1)

    # ============ G. ASESMEN PEMBELAJARAN ============
    _heading(doc, "G. ASESMEN PEMBELAJARAN", 3)

    _heading(doc, "ASESMEN DIAGNOSTIK", 4)
    for item in _g(ai_asesmen, 'asesmen_diagnostik', ['-']):
        add_plain_bullet(doc, item)

    _heading(doc, "ASESMEN FORMATIF", 4)
    for item in _g(ai_asesmen, 'asesmen_formatif', ['-']):
        add_plain_bullet(doc, item)

    _heading(doc, "ASESMEN SUMATIF", 4)
    for item in _g(ai_asesmen, 'asesmen_sumatif', ['-']):
        add_plain_bullet(doc, item)

    add_body_paragraph(doc, "Contoh Tes Tertulis :", bold=True, italic=True)
    add_body_paragraph(doc, "Pilihan Ganda", bold=True)
    for i, soal in enumerate(_g(ai_asesmen, 'soal_pg', [])):
        add_manual_numbered(doc, i + 1, soal.get('soal', '-'))
        for opt in ('a', 'b', 'c', 'd'):
            add_plain_bullet(doc, f"{opt}. {soal.get(opt, '-')}", level=1)

    add_body_paragraph(doc, "Esai", bold=True)
    for i, soal in enumerate(_g(ai_asesmen, 'soal_esai', ['-'])):
        add_manual_numbered(doc, i + 1, soal)

    # ---------------- TANDA TANGAN (tabel — format dipertahankan) ----------------
    doc.add_paragraph()
    ttd_table = doc.add_table(rows=1, cols=2)
    ttd_table.autofit = False
    for cell in ttd_table.rows[0].cells:
        cell.width = Cm(8.25)
    c1 = ttd_table.cell(0, 0)
    c1.text = f"Mengetahui,\nKepala Sekolah\n\n\n\n( {data_input['kepsek']} )"
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2 = ttd_table.cell(0, 1)
    c2.text = f"{data_input['sekolah']}, {data_input.get('tahun_pelajaran', '')}\nGuru Mata Pelajaran\n\n\n\n( {data_input['guru']} )"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
