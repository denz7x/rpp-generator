"""
Helper functions untuk membangun dokumen .docx bergaya
'MODUL AJAR KURIKULUM MERDEKA (Deep Learning)' resmi Kemendikbud,
menggunakan python-docx.

Semua fungsi tabel di bawah ini (set_cell_background, set_table_borders,
set_col_widths, ModulTable, dst) TIDAK DIUBAH dari versi sebelumnya, supaya
format tabel (cover & tanda tangan) yang sudah rapi tetap identik.
Hanya ditambahkan helper baru di bagian paling bawah untuk kebutuhan bullet
berlabel tebal (dipakai oleh format Deep Learning yang lebih banyak berbasis
heading + bullet, bukan tabel besar).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_FILL = "D9D9D9"   # abu-abu, sama seperti dokumen contoh
WHITE_FILL = "FFFFFF"
FONT_NAME = "Times New Roman"
FONT_SIZE = 11


# ----------------------------------------------------------------------
# UTILITAS TINGKAT RENDAH (OXML) — TIDAK DIUBAH
# ----------------------------------------------------------------------

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_borders(table, size=4, color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, widths_cm):
    """widths_cm: list lebar kolom dalam cm. Perlu diset di setiap row/cell."""
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def add_bottom_border(paragraph, size=12, color="000000"):
    """Menambahkan garis horizontal (border bawah) pada sebuah paragraf kosong."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_repeat_table_header(row):
    """Menandai baris sebagai header yang berulang di setiap halaman."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def style_run(run, bold=False, italic=False, size=FONT_SIZE, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # Pastikan font timur (east asian) juga ikut, biar konsisten di semua platform
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_paragraph_in_cell(cell, text="", bold=False, italic=False, size=FONT_SIZE,
                           align=None, space_after=4, first=False, bullet_style=None):
    """Tambahkan paragraf ke dalam cell tabel. Jika first=True, pakai paragraf
    kosong pertama yang sudah ada di cell alih-alih membuat baru."""
    if first and len(cell.paragraphs) == 1 and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    if bullet_style:
        p.style = bullet_style
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        style_run(run, bold=bold, italic=italic, size=size)
    return p


# ----------------------------------------------------------------------
# BUILDER TABEL UTAMA — TIDAK DIUBAH (disisakan untuk kompatibilitas /
# dipakai ulang bila suatu saat body kembali membutuhkan tabel besar)
# ----------------------------------------------------------------------

class ModulTable:
    def __init__(self, doc, col_widths_cm=(4.5, 12.0)):
        self.doc = doc
        self.col_widths_cm = col_widths_cm
        self.table = doc.add_table(rows=1, cols=2)
        self.table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(self.table)
        set_col_widths(self.table, col_widths_cm)
        self._first_row_used = False

    def _new_row(self):
        if not self._first_row_used:
            row = self.table.rows[0]
            self._first_row_used = True
        else:
            row = self.table.add_row()
        set_col_widths(self.table, self.col_widths_cm)
        return row

    def add_section_header(self, text, shade=HEADER_FILL, size=12):
        row = self._new_row()
        merged = row.cells[0].merge(row.cells[1])
        set_cell_background(merged, shade)
        merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        merged.paragraphs[0].text = ""
        add_paragraph_in_cell(merged, text, bold=True, size=size, first=True)
        return merged

    def add_label_value(self, label, value):
        row = self._new_row()
        c0, c1 = row.cells
        c0.paragraphs[0].text = ""
        add_paragraph_in_cell(c0, label, bold=True, first=True)
        c1.paragraphs[0].text = ""
        add_paragraph_in_cell(c1, f": {value}" if value else ": ", first=True)
        return row

    def add_full_content(self, content, bullet=False, italic_intro=None):
        row = self._new_row()
        merged = row.cells[0].merge(row.cells[1])
        merged.paragraphs[0].text = ""
        first = True

        if italic_intro:
            add_paragraph_in_cell(merged, italic_intro, italic=True, first=first)
            first = False

        if content is None:
            items = []
        else:
            items = content if isinstance(content, list) else [content]
        if not items and not italic_intro:
            items = ["-"]

        for item in items:
            style = 'List Bullet' if bullet else None
            add_paragraph_in_cell(merged, str(item), first=first, bullet_style=style)
            first = False
        return merged

    def add_nested_table(self, headers, rows, col_widths_cm=None, merge_first_row_span=None):
        row = self._new_row()
        merged = row.cells[0].merge(row.cells[1])
        merged.paragraphs[0].text = ""

        n_cols = len(headers)
        inner = merged.add_table(rows=1, cols=n_cols)
        empty_p = merged.paragraphs[0]._p
        empty_p.getparent().remove(empty_p)
        set_table_borders(inner, size=3)
        if col_widths_cm:
            set_col_widths(inner, col_widths_cm)

        hdr_cells = inner.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].paragraphs[0].text = ""
            set_cell_background(hdr_cells[i], HEADER_FILL)
            add_paragraph_in_cell(hdr_cells[i], h, bold=True, size=10, first=True,
                                   align=WD_ALIGN_PARAGRAPH.CENTER)
        set_repeat_table_header(inner.rows[0])

        for r in rows:
            data_row = inner.add_row()
            for i, val in enumerate(r):
                cell = data_row.cells[i]
                cell.paragraphs[0].text = ""
                add_paragraph_in_cell(cell, str(val), size=10, first=True)

        if col_widths_cm:
            set_col_widths(inner, col_widths_cm)
        return inner


# ----------------------------------------------------------------------
# HELPER BARU: bullet dengan label tebal di depan, mengikuti gaya dokumen
# contoh "Modul Ajar Deep Learning" (mis. "Salam dan Doa: Guru membuka...")
# ----------------------------------------------------------------------

def add_bold_lead_bullet(doc, lead, rest, level=0, size=FONT_SIZE):
    """
    Tambahkan bullet ke body dokumen (bukan cell tabel) dengan format:
    "**lead**: rest". level=0 -> 'List Bullet', level=1 -> 'List Bullet 2'.
    Jika lead kosong, seluruh teks 'rest' ditulis bold (dipakai untuk
    sub-judul semacam "KEGIATAN INTI (55 MENIT)").
    """
    style = 'List Bullet 2' if level >= 1 else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    if lead:
        r1 = p.add_run(f"{lead}: ")
        style_run(r1, bold=True, size=size)
        r2 = p.add_run(str(rest))
        style_run(r2, bold=False, size=size)
    else:
        r = p.add_run(str(rest))
        style_run(r, bold=True, size=size)
    return p


def add_plain_bullet(doc, text, level=0, italic=False, size=FONT_SIZE):
    style = 'List Bullet 2' if level >= 1 else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(str(text))
    style_run(r, italic=italic, size=size)
    return p


def add_numbered_item(doc, text, size=FONT_SIZE):
    """Item bernomor. Style bawaan Word 'List Number' melanjutkan penomoran
    lintas-section (tidak reset), jadi di sini nomor ditulis manual lewat
    parameter number pada add_manual_numbered."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(str(text))
    style_run(r, size=size)
    return p


def add_manual_numbered(doc, number, text, size=FONT_SIZE, indent_cm=0.63):
    """Item bernomor manual ('1.  teks'), tiap pemanggilan section baru bisa
    mulai dari nomor 1 lagi tanpa terpengaruh nomor section sebelumnya."""
    from docx.shared import Cm as _Cm
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = _Cm(indent_cm)
    p.paragraph_format.first_line_indent = _Cm(-indent_cm)
    r = p.add_run(f"{number}.\t{text}")
    style_run(r, size=size)
    return p


def add_body_paragraph(doc, text, bold=False, italic=False, size=FONT_SIZE,
                        align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    style_run(r, bold=bold, italic=italic, size=size)
    return p


def set_heading_font(paragraph, size=None):
    """Paksa heading tetap pakai Times New Roman (default Word heading pakai
    font tema lain) supaya konsisten dengan seluruh dokumen."""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        if size:
            run.font.size = Pt(size)
