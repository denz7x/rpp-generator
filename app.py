import streamlit as st
import google.generativeai as genai
import json
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 1. KONFIGURASI API KEY
# ==========================================
try:
    MY_API_KEY = st.secrets["GOOGLE_API_KEY"]
except Exception:
    MY_API_KEY = None
    st.error(
        "⚠️ API Key belum diatur. Tambahkan GOOGLE_API_KEY di menu "
        "Settings → Secrets (Streamlit Cloud), atau di file "
        ".streamlit/secrets.toml (kalau dijalankan lokal)."
    )
    st.stop()

# Konfigurasi Awal
try:
    genai.configure(api_key=MY_API_KEY)
except Exception as e:
    st.error(f"Error Konfigurasi: {e}")
    st.stop()

# ==========================================
# 2. PENGATURAN TAMPILAN (UI) COLORFUL & ESTETIK
# ==========================================
st.set_page_config(
    page_title="EduGen Pro - RPP Generator",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS Super Modern, Colorful & Glassmorphism
st.markdown("""
<style>
    /* Global Background: Lembut dengan sentuhan pastel */
    .stApp {
        background-color: #f4f7f6;
        background-image: radial-gradient(circle at 10% 20%, rgba(255, 117, 140, 0.05) 0%, transparent 40%),
                          radial-gradient(circle at 90% 80%, rgba(0, 242, 254, 0.05) 0%, transparent 40%);
    }

    /* Container Utama */
    .main {
        padding: 1rem 2rem;
    }

    /* Header Super Estetik */
    .header-container {
        background: linear-gradient(135deg, #ff758c 0%, #ff7eb3 100%);
        padding: 2.5rem 2rem;
        border-radius: 24px;
        margin-bottom: 2.5rem;
        color: white;
        text-align: center;
        box-shadow: 0 10px 30px rgba(255, 117, 140, 0.3);
        position: relative;
        overflow: hidden;
    }
    
    /* Efek kilauan di header */
    .header-container::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255,255,255,0.2) 0%, rgba(255,255,255,0) 60%);
        transform: rotate(30deg);
        pointer-events: none;
    }

    /* Card Melayang (Soft UI) */
    .stCard {
        background: rgba(255, 255, 255, 0.85);
        backdrop-filter: blur(12px);
        padding: 2rem;
        border-radius: 20px;
        box-shadow: 0 8px 32px rgba(31, 38, 135, 0.05);
        margin-bottom: 1.5rem;
        border: 1px solid rgba(255, 255, 255, 0.4);
        border-top: 5px solid #00f2fe;
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }
    .stCard:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 40px rgba(31, 38, 135, 0.1);
    }

    /* Tombol Standar (Cyan/Blue) */
    .stButton>button {
        background: linear-gradient(45deg, #4facfe 0%, #00f2fe 100%);
        color: white;
        border: none;
        padding: 0.6rem 2rem;
        border-radius: 50px;
        font-weight: 700;
        letter-spacing: 0.5px;
        transition: all 0.3s ease;
        box-shadow: 0 8px 15px rgba(0, 242, 254, 0.3);
        width: 100%;
        margin-top: 1rem;
    }
    .stButton>button:hover {
        transform: translateY(-3px) scale(1.02);
        box-shadow: 0 12px 20px rgba(0, 242, 254, 0.4);
        background: linear-gradient(45deg, #00f2fe 0%, #4facfe 100%);
    }

    /* Tombol Primary Keras (Merah/Peach untuk Generate) */
    button[kind="primary"] {
        background: linear-gradient(45deg, #ff0844 0%, #ffb199 100%) !important;
        box-shadow: 0 8px 15px rgba(255, 8, 68, 0.3) !important;
        font-size: 1.1rem !important;
        padding: 0.8rem 2rem !important;
    }
    button[kind="primary"]:hover {
        background: linear-gradient(45deg, #ffb199 0%, #ff0844 100%) !important;
        box-shadow: 0 12px 25px rgba(255, 8, 68, 0.4) !important;
    }

    /* Input Fields Luwes */
    .stTextInput>div>div>input, .stSelectbox>div>div>select, .stMultiSelect>div>div>div {
        border-radius: 12px;
        border: 2px solid #e2e8f0;
        padding: 0.6rem 1rem;
        background-color: #f8fafc;
        transition: all 0.3s ease;
    }
    .stTextInput>div>div>input:focus, .stSelectbox>div>div>select:focus {
        border-color: #ff7eb3;
        box-shadow: 0 0 0 3px rgba(255, 126, 179, 0.2);
        background-color: #ffffff;
    }

    /* Tabs Styling (Estetik) */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
        background-color: rgba(255,255,255,0.6);
        padding: 0.5rem;
        border-radius: 50px;
        backdrop-filter: blur(5px);
        margin-bottom: 1rem;
    }
    .stTabs [data-baseweb="tab"] {
        height: 45px;
        background-color: transparent;
        border-radius: 50px;
        color: #64748b;
        font-weight: 600;
        border: none;
        padding: 0 1.5rem;
        transition: all 0.3s ease;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        box-shadow: 0 4px 10px rgba(102, 126, 234, 0.3);
    }

    /* Pesan Sukses Vibrant */
    .success-message {
        background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
        color: white;
        padding: 1rem 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
        font-weight: bold;
        box-shadow: 0 4px 15px rgba(56, 239, 125, 0.3);
        display: flex;
        align-items: center;
        gap: 10px;
    }

    /* Sidebar Sidebar Keren */
    [data-testid="stSidebar"] {
        background-image: linear-gradient(180deg, #ffffff 0%, #fef2f2 100%);
        border-right: 1px solid #fce7f3;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. DATABASE (SESSION STATE)
# ==========================================
if 'profil_db' not in st.session_state:
    st.session_state['profil_db'] = [
        "Beriman, Bertakwa kepada Tuhan YME, dan Berakhlak Mulia",
        "Berkebinekaan Global",
        "Bergotong Royong",
        "Mandiri",
        "Bernalar Kritis",
        "Kreatif"
    ]

if 'ai_result' not in st.session_state:
    st.session_state.ai_result = None

# ==========================================
# 4. FUNGSI LOGIKA (BACKEND)
# ==========================================
def get_available_model():
    try:
        available_models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                available_models.append(m.name)
        
        if not available_models: return None
        prioritas = [
            "models/gemini-2.5-flash",
            "models/gemini-1.5-flash",
        ]
        for nama in prioritas:
            if nama in available_models:
                return nama
        return available_models[0]
    except Exception as e:
        st.error(f"Gagal mengambil daftar model dari Google: {e}")
        return None

def generate_rpp_content(model_name, mapel, topik, kelas, waktu, profil_list, pakai_lkpd):
    try:
        model = genai.GenerativeModel(model_name)
        profil_str = ", ".join(profil_list)
        
        instruksi_lkpd = ""
        json_structure_lkpd = ""
        if pakai_lkpd == "Ya":
            instruksi_lkpd = "Sertakan juga materi untuk Lembar Kerja Peserta Didik (LKPD) berisi 3-5 soal atau aktivitas."
            json_structure_lkpd = ', "lkpd": "Isi detail LKPD (Soal/Aktivitas)."'

        prompt = f"""
        Buatkan Modul Ajar Kurikulum Merdeka dalam format JSON.
        Data: Mapel {mapel}, Kelas {kelas}, Topik {topik}, Waktu {waktu}, Profil {profil_str}.
        {instruksi_lkpd}
        
        Output WAJIB JSON MURNI (tanpa format markdown):
        {{
            "tujuan": "Tujuan pembelajaran (poin-poin).",
            "pemahaman": "Pertanyaan pemantik.",
            "pendahuluan": "Kegiatan awal (poin-poin).",
            "inti": "Kegiatan inti detail (poin-poin).",
            "penutup": "Kegiatan penutup (poin-poin).",
            "asesmen": "Teknik penilaian."
            {json_structure_lkpd}
        }}
        Gunakan Bahasa Indonesia formal pendidikan.
        """
        
        response = model.generate_content(
            prompt,
            request_options={"timeout": 60},
        )
        text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except json.JSONDecodeError:
        st.error("⚠️ AI mengembalikan format yang tidak sesuai. Coba klik tombol Generate sekali lagi.")
        return None
    except Exception as e:
        st.error(f"Gagal Generate: {str(e)}")
        return None

def create_docx(data_input, ai_data, pakai_lkpd):
    doc = Document()
    
    head = doc.add_heading('MODUL AJAR / RPP', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(0.2)
    table.columns[2].width = Inches(4.5)
    
    infos = [
        ("Nama Sekolah", data_input['sekolah']),
        ("Nama Guru", data_input['guru']),
        ("Mata Pelajaran", data_input['mapel']),
        ("Kelas / Semester", data_input['kelas']),
        ("Alokasi Waktu", data_input['waktu'])
    ]
    
    for i, (label, val) in enumerate(infos):
        table.cell(i,0).text = label
        table.cell(i,1).text = ":"
        table.cell(i,2).text = val
        table.cell(i,0).paragraphs[0].paragraph_format.space_after = Pt(2)
        table.cell(i,2).paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")

    def add_section(title, content):
        doc.add_heading(title, level=1)
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_paragraph("-")

    add_section('A. Tujuan Pembelajaran', ai_data.get('tujuan'))
    
    doc.add_heading('B. Profil Pelajar Pancasila', level=1)
    for p in data_input['profil']:
        doc.add_paragraph(f"- {p}", style='List Bullet')

    add_section('C. Pemahaman Bermakna', ai_data.get('pemahaman'))
    
    doc.add_heading('D. Kegiatan Pembelajaran', level=1)
    
    p = doc.add_paragraph()
    p.add_run("1. Kegiatan Pendahuluan").bold = True
    doc.add_paragraph(ai_data.get('pendahuluan', '-'))
    
    p = doc.add_paragraph()
    p.add_run("2. Kegiatan Inti").bold = True
    doc.add_paragraph(ai_data.get('inti', '-'))
    
    p = doc.add_paragraph()
    p.add_run("3. Kegiatan Penutup").bold = True
    doc.add_paragraph(ai_data.get('penutup', '-'))

    add_section('E. Asesmen / Penilaian', ai_data.get('asesmen'))

    doc.add_paragraph("\n\n")
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    
    c1 = sig_table.cell(0,0)
    c1.text = f"Mengetahui,\nKepala Sekolah\n\n\n\n{data_input['kepsek']}"
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c2 = sig_table.cell(0,1)
    c2.text = f"Guru Mata Pelajaran\n\n\n\n{data_input['guru']}"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if pakai_lkpd == "Ya" and ai_data.get('lkpd'):
        doc.add_page_break()
        doc.add_heading('LAMPIRAN: LEMBAR KERJA PESERTA DIDIK (LKPD)', 0)
        doc.add_paragraph("")
        doc.add_paragraph("Nama Siswa : ...................................")
        doc.add_paragraph(f"Kelas      : {data_input['kelas']}")
        doc.add_paragraph("----------------------------------------------------------------------------------")
        doc.add_paragraph(ai_data.get('lkpd'))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 5. HALAMAN UTAMA
# ==========================================
def page_generator():
    st.markdown("""
    <div class="header-container">
        <h1 style="margin: 0; font-size: 3rem; text-shadow: 2px 2px 4px rgba(0,0,0,0.1);">✨ EduGen Pro</h1>
        <p style="margin: 0.5rem 0 0 0; font-size: 1.2rem; font-weight: 500; opacity: 0.95;">
            Sistem Cerdas Penyusun Modul Ajar & LKPD
        </p>
        <div style="margin-top: 15px; display: inline-block; background: rgba(255,255,255,0.25); padding: 6px 18px; border-radius: 50px; font-size: 0.95rem; backdrop-filter: blur(5px); font-weight: bold;">
            Ceng Ucu Muhammad, S.H - SMP IT Nurusy Syifa
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    active_model = get_available_model()
    if not active_model:
        st.error("⚠️ API Key bermasalah atau kuota habis.")
        st.stop()

    tab1, tab2, tab3 = st.tabs(["📝 Input Data", "👁️ Preview Hasil", "⚙️ Pengaturan"])
    
    with tab1:
        with st.container():
            st.markdown('<div class="stCard">', unsafe_allow_html=True)
            st.markdown("<h3 style='color: #334155;'>1. Identitas Sekolah & Guru</h3>", unsafe_allow_html=True)
            col1, col2, col3 = st.columns(3)
            with col1:
                nama_guru = st.text_input("Nama Guru", placeholder="Cth: Budi Santoso, S.Pd")
            with col2:
                nama_sekolah = st.text_input("Nama Sekolah", value="SMP IT Nurusy Syifa")
            with col3:
                nama_kepsek = st.text_input("Nama Kepala Sekolah", placeholder="Cth: Ahmad, M.Pd")
            st.markdown('</div>', unsafe_allow_html=True)
        
        with st.container():
            st.markdown('<div class="stCard">', unsafe_allow_html=True)
            st.markdown("<h3 style='color: #334155;'>2. Parameter Materi</h3>", unsafe_allow_html=True)
            col4, col5 = st.columns(2)
            with col4:
                mapel = st.text_input("Mata Pelajaran", value="Ilmu Pengetahuan Sosial (IPS)")
                kelas = st.selectbox("Kelas", ["VII (Fase D)", "VIII (Fase D)", "IX (Fase D)"], index=0)
                waktu = st.text_input("Alokasi Waktu", value="2 JP (2 x 40 Menit)")
            with col5:
                topik = st.text_input("Topik Materi*", placeholder="Wajib diisi. Cth: Interaksi Sosial")
                profil = st.multiselect(
                    "Profil Pelajar Pancasila",
                    st.session_state['profil_db'],
                    default=st.session_state['profil_db'][:2]
                )
                pilihan_lkpd = st.radio(
                    "Sertakan Lembar Kerja (LKPD)?",
                    ["Tidak", "Ya"],
                    horizontal=True
                )
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
        with col_btn2:
            submitted = st.button("🚀 GENERATE MODUL SEKARANG", use_container_width=True, type="primary")
    
    with tab2:
        if st.session_state.ai_result:
            st.markdown('<div class="success-message">🎉 Sempurna! Modul Ajar telah selesai disusun.</div>', unsafe_allow_html=True)
            with st.expander("🎯 Tujuan Pembelajaran", expanded=True):
                st.write(st.session_state.ai_result.get('tujuan', 'Tidak tersedia'))
            
            col_preview1, col_preview2 = st.columns(2)
            with col_preview1:
                with st.expander("🔥 Kegiatan Inti", expanded=True):
                    st.write(st.session_state.ai_result.get('inti', 'Tidak tersedia'))
            with col_preview2:
                with st.expander("📝 Asesmen / Penilaian", expanded=True):
                    st.write(st.session_state.ai_result.get('asesmen', 'Tidak tersedia'))
            
            if 'lkpd' in st.session_state.ai_result:
                with st.expander("📚 Lembar Kerja (LKPD)", expanded=True):
                    st.info(st.session_state.ai_result.get('lkpd'))
        else:
            st.info("💡 Isi parameter di tab 'Input Data' lalu klik Generate untuk melihat keajaiban AI di sini.")
    
    with tab3:
        with st.container():
            st.markdown('<div class="stCard">', unsafe_allow_html=True)
            st.markdown("<h3 style='color: #334155;'>Informasi Sistem</h3>", unsafe_allow_html=True)
            st.info(f"🧠 Mesin AI Aktif: **{active_model.split('/')[-1]}**")
            st.markdown("""
            **Lisensi & Hak Cipta:**  
            Aplikasi khusus internal yang dikembangkan secara mandiri untuk menunjang administrasi digital guru di lingkungan sekolah.
            """)
            st.markdown('</div>', unsafe_allow_html=True)
    
    if 'submitted' in locals() and submitted:
        if not topik:
            st.error("⚠️ Harap isi 'Topik Materi' terlebih dahulu!")
        elif not nama_guru or not nama_sekolah:
            st.error("⚠️ Nama Guru dan Sekolah tidak boleh kosong!")
        else:
            with st.spinner("✨ AI sedang meracik Modul Ajar & LKPD... Harap bersabar ya!"):
                res = generate_rpp_content(active_model, mapel, topik, kelas, waktu, profil, pilihan_lkpd)
                if res:
                    st.session_state.ai_result = res
                    st.session_state.data_input = {
                        'guru': nama_guru, 'sekolah': nama_sekolah, 'kepsek': nama_kepsek,
                        'mapel': mapel, 'kelas': kelas, 'waktu': waktu, 'profil': profil,
                        'pilihan_lkpd': pilihan_lkpd
                    }
                    st.success("✅ Modul berhasil dibuat! Silakan pindah ke tab 'Preview Hasil'.")
                    st.rerun()

    if st.session_state.ai_result and st.session_state.get('data_input'):
        st.markdown("<br><hr>", unsafe_allow_html=True)
        col_dl1, col_dl2, col_dl3 = st.columns([1, 2, 1])
        with col_dl2:
            docx_file = create_docx(
                st.session_state.data_input, 
                st.session_state.ai_result, 
                st.session_state.data_input['pilihan_lkpd']
            )
            st.download_button(
                label="📥 UNDUH MODUL AJAR (.DOCX)",
                data=docx_file,
                file_name=f"Modul_Ajar_{st.session_state.data_input['mapel']}_{st.session_state.data_input['kelas']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

# ==========================================
# 6. HALAMAN DATABASE PROFIL
# ==========================================
def page_profil():
    st.markdown("""
    <div style="background: linear-gradient(135deg, #a18cd1 0%, #fbc2eb 100%); padding: 1.5rem; border-radius: 20px; color: white; margin-bottom: 2rem; box-shadow: 0 4px 15px rgba(161, 140, 209, 0.3);">
        <h2 style="margin:0;">🎓 Database Profil Pelajar Pancasila</h2>
    </div>
    """, unsafe_allow_html=True)
    
    with st.container():
        st.markdown('<div class="stCard">', unsafe_allow_html=True)
        st.markdown("<h4 style='color: #334155;'>➕ Tambah Profil Baru</h4>", unsafe_allow_html=True)
        col_add1, col_add2 = st.columns([3, 1])
        with col_add1:
            baru = st.text_input("Nama profil", placeholder="Cth: Mandiri Berwawasan Lingkungan", label_visibility="collapsed")
        with col_add2:
            if st.button("Simpan", use_container_width=True) and baru:
                if baru not in st.session_state['profil_db']:
                    st.session_state['profil_db'].append(baru)
                    st.success(f"Tersimpan: {baru}")
                    st.rerun()
                else:
                    st.warning("Profil sudah ada.")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with st.container():
        st.markdown('<div class="stCard">', unsafe_allow_html=True)
        st.markdown("<h4 style='color: #334155;'>📋 Daftar Profil Aktif</h4>", unsafe_allow_html=True)
        for i, p in enumerate(st.session_state['profil_db']):
            col_prof1, col_prof2 = st.columns([4, 1])
            with col_prof1:
                st.markdown(f"<div style='padding: 8px; background: #f8fafc; border-radius: 8px; border: 1px solid #e2e8f0; margin-bottom: 5px;'><b>{i+1}.</b> {p}</div>", unsafe_allow_html=True)
            with col_prof2:
                if st.button("Hapus", key=f"del_{i}", use_container_width=True):
                    st.session_state['profil_db'].pop(i)
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# ==========================================
# 7. HALAMAN TENTANG
# ==========================================
def page_tentang():
    st.markdown("""
    <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 1.5rem; border-radius: 20px; color: white; margin-bottom: 2rem; box-shadow: 0 4px 15px rgba(0, 242, 254, 0.3);">
        <h2 style="margin:0;">ℹ️ Tentang Aplikasi</h2>
    </div>
    """, unsafe_allow_html=True)
    
    with st.container():
        st.markdown('<div class="stCard">', unsafe_allow_html=True)
        st.markdown("""
        ### EduGen Pro (v2.0.0)
        
        Sebuah alat bantu pintar *(Smart Tools)* berbasis kecerdasan buatan (AI) yang dirombak khusus dengan antarmuka modern guna mempermudah administrasi pengajaran.
        
        **Pengembang:** Ceng Ucu Muhammad, S.H  
        **Instansi:** SMP IT Nurusy Syifa  
        
        *Didesain dengan cinta untuk kemajuan pendidikan Indonesia.*
        """)
        st.markdown('</div>', unsafe_allow_html=True)

# ==========================================
# 8. NAVIGASI UTAMA
# ==========================================
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding: 1.5rem 0; background: linear-gradient(135deg, #ff758c 0%, #ff7eb3 100%); border-radius: 16px; color: white; margin-bottom: 1.5rem; box-shadow: 0 4px 15px rgba(255, 117, 140, 0.3);">
        <h2 style="margin: 0; font-size: 2rem; color: white; text-shadow: 1px 1px 2px rgba(0,0,0,0.1);">🌟 EduGen</h2>
        <p style="margin: 0; font-size: 0.9rem; opacity: 0.95; font-weight: 500;">SMP IT Nurusy Syifa</p>
    </div>
    """, unsafe_allow_html=True)
    
    menu_options = {
        "📝 Buat Modul Ajar": page_generator,
        "🎓 Kelola Profil": page_profil,
        "ℹ️ Tentang": page_tentang
    }
    
    menu_selection = st.radio(
        "Navigasi:",
        list(menu_options.keys()),
        label_visibility="collapsed"
    )
    
    st.markdown("<br><hr>", unsafe_allow_html=True)
    if st.button("🔄 Bersihkan Sesi", use_container_width=True):
        for key in list(st.session_state.keys()):
            if key != 'profil_db':
                del st.session_state[key]
        st.rerun()

menu_options[menu_selection]()
